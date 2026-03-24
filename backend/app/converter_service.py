"""Document conversion service with support for multiple formats."""

import logging
import os
import tempfile
from pathlib import Path
from io import BytesIO

import cv2
import easyocr
import fitz  # PyMuPDF
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from openpyxl import Workbook, load_workbook
from PIL import Image
from pdf2docx import Converter as Pdf2DocxConverter

# doctr for advanced OCR with layout analysis
from doctr.io import DocumentFile
from doctr.models import ocr_predictor

from app.converter_schemas import ConversionType

# Setup logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# Initialize EasyOCR reader (lazy loaded)
_ocr_reader = None

# Initialize doctr predictor (lazy loaded)
_doctr_predictor = None

def get_ocr_reader():
    """Get or create OCR reader (singleton)."""
    global _ocr_reader
    if _ocr_reader is None:
        _ocr_reader = easyocr.Reader(['en'], gpu=False)
    return _ocr_reader


def get_doctr_predictor():
    """Get or create doctr OCR predictor (singleton)."""
    global _doctr_predictor
    if _doctr_predictor is None:
        logger.info("Initializing doctr OCR predictor...")
        _doctr_predictor = ocr_predictor(pretrained=True)
        logger.info("doctr predictor initialized")
    return _doctr_predictor


def add_floating_textbox(paragraph, text: str, x_emu: int, y_emu: int,
                         width_emu: int, height_emu: int, font_size_pt: float = 11):
    """Add a positioned text box (shape) at exact coordinates in Word.

    Uses DrawingML to create an absolutely positioned text box.
    Coordinates are in EMUs (914400 EMUs = 1 inch).
    """
    # Create the inline shape container
    run = paragraph.add_run()

    # Create drawing element
    drawing = OxmlElement('w:drawing')

    # Create anchor for absolute positioning
    anchor = OxmlElement('wp:anchor')
    anchor.set(qn('wp:distT'), '0')
    anchor.set(qn('wp:distB'), '0')
    anchor.set(qn('wp:distL'), '0')
    anchor.set(qn('wp:distR'), '0')
    anchor.set(qn('wp:simplePos'), '0')
    anchor.set(qn('wp:relativeHeight'), '0')
    anchor.set(qn('wp:behindDoc'), '0')
    anchor.set(qn('wp:locked'), '0')
    anchor.set(qn('wp:layoutInCell'), '1')
    anchor.set(qn('wp:allowOverlap'), '1')

    # Simple position (not used but required)
    simplePos = OxmlElement('wp:simplePos')
    simplePos.set('x', '0')
    simplePos.set('y', '0')
    anchor.append(simplePos)

    # Horizontal position - absolute from page
    posH = OxmlElement('wp:positionH')
    posH.set('relativeFrom', 'page')
    posOffset_h = OxmlElement('wp:posOffset')
    posOffset_h.text = str(x_emu)
    posH.append(posOffset_h)
    anchor.append(posH)

    # Vertical position - absolute from page
    posV = OxmlElement('wp:positionV')
    posV.set('relativeFrom', 'page')
    posOffset_v = OxmlElement('wp:posOffset')
    posOffset_v.text = str(y_emu)
    posV.append(posOffset_v)
    anchor.append(posV)

    # Extent (size)
    extent = OxmlElement('wp:extent')
    extent.set('cx', str(width_emu))
    extent.set('cy', str(height_emu))
    anchor.append(extent)

    # Effect extent
    effectExtent = OxmlElement('wp:effectExtent')
    effectExtent.set('l', '0')
    effectExtent.set('t', '0')
    effectExtent.set('r', '0')
    effectExtent.set('b', '0')
    anchor.append(effectExtent)

    # Wrap none (text flows around)
    wrapNone = OxmlElement('wp:wrapNone')
    anchor.append(wrapNone)

    # Doc properties
    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(id(text) % 10000))
    docPr.set('name', 'TextBox')
    anchor.append(docPr)

    # Graphic frame
    graphic = OxmlElement('a:graphic')
    graphic.set(qn('xmlns:a'), 'http://schemas.openxmlformats.org/drawingml/2006/main')

    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape')

    # Word processing shape
    wsp = OxmlElement('wps:wsp')

    # Shape properties
    spPr = OxmlElement('wps:spPr')

    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext = OxmlElement('a:ext')
    ext.set('cx', str(width_emu))
    ext.set('cy', str(height_emu))
    xfrm.append(ext)
    spPr.append(xfrm)

    # Rectangle preset
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)

    # No fill (transparent)
    noFill = OxmlElement('a:noFill')
    spPr.append(noFill)

    # No line (no border)
    ln = OxmlElement('a:ln')
    noFill2 = OxmlElement('a:noFill')
    ln.append(noFill2)
    spPr.append(ln)

    wsp.append(spPr)

    # Text box content
    txbx = OxmlElement('wps:txbx')
    txbxContent = OxmlElement('w:txbxContent')

    p = OxmlElement('w:p')
    r = OxmlElement('w:r')

    # Font size
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size_pt * 2)))  # Half-points
    rPr.append(sz)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    txbxContent.append(p)
    txbx.append(txbxContent)
    wsp.append(txbx)

    # Body properties
    bodyPr = OxmlElement('wps:bodyPr')
    bodyPr.set('rot', '0')
    bodyPr.set('vert', 'horz')
    bodyPr.set('wrap', 'square')
    bodyPr.set('lIns', '0')
    bodyPr.set('tIns', '0')
    bodyPr.set('rIns', '0')
    bodyPr.set('bIns', '0')
    bodyPr.set('anchor', 't')
    bodyPr.set('anchorCtr', '0')
    wsp.append(bodyPr)

    graphicData.append(wsp)
    graphic.append(graphicData)
    anchor.append(graphic)
    drawing.append(anchor)

    run._r.append(drawing)


def detect_columns_from_words(words, min_gap_ratio=0.03):
    """Detect column boundaries from word positions based on gaps.

    Args:
        words: List of word dicts with x_min, x_max positions (normalized 0-1)
        min_gap_ratio: Minimum gap size as ratio of page width to consider as column separator

    Returns:
        List of column boundaries (x positions) or None if no clear columns
    """
    if len(words) < 5:
        return None

    # Get all x positions (start and end of each word)
    x_positions = []
    for w in words:
        x_positions.append(('start', w['x_min']))
        x_positions.append(('end', w['x_max']))

    # Sort by position
    x_positions.sort(key=lambda x: x[1])

    # Find gaps between word ends and next word starts
    gaps = []
    i = 0
    while i < len(x_positions) - 1:
        if x_positions[i][0] == 'end' and x_positions[i+1][0] == 'start':
            gap_size = x_positions[i+1][1] - x_positions[i][1]
            if gap_size > min_gap_ratio:
                gaps.append({
                    'start': x_positions[i][1],
                    'end': x_positions[i+1][1],
                    'center': (x_positions[i][1] + x_positions[i+1][1]) / 2,
                    'size': gap_size
                })
        i += 1

    if len(gaps) < 2:
        return None

    # Cluster gaps by x position to find consistent column separators
    gap_clusters = []
    gaps.sort(key=lambda g: g['center'])

    current_cluster = [gaps[0]]
    for gap in gaps[1:]:
        if gap['center'] - current_cluster[-1]['center'] < 0.05:  # 5% tolerance
            current_cluster.append(gap)
        else:
            if len(current_cluster) >= 2:  # Need at least 2 gaps at similar position
                gap_clusters.append(current_cluster)
            current_cluster = [gap]

    if len(current_cluster) >= 2:
        gap_clusters.append(current_cluster)

    if len(gap_clusters) < 2:  # Need at least 2 column separators for 3+ columns
        return None

    # Get column boundaries from gap clusters
    col_boundaries = [0.0]  # Start of page
    for cluster in gap_clusters:
        avg_center = sum(g['center'] for g in cluster) / len(cluster)
        col_boundaries.append(avg_center)
    col_boundaries.append(1.0)  # End of page

    return col_boundaries


def detect_table_rows_from_words(words, line_threshold=0.015):
    """Group words into rows based on Y position.

    Returns list of rows, each row is a list of words.
    """
    if not words:
        return []

    # Sort by Y position
    sorted_words = sorted(words, key=lambda w: w['y_min'])

    rows = []
    current_row = [sorted_words[0]]
    current_y = sorted_words[0]['y_min']

    for word in sorted_words[1:]:
        if abs(word['y_min'] - current_y) < line_threshold:
            current_row.append(word)
        else:
            rows.append(current_row)
            current_row = [word]
            current_y = word['y_min']

    if current_row:
        rows.append(current_row)

    return rows


def is_footer_content(row):
    """Check if row is footer content (bottom of page).

    Footer typically contains 'Powered by' and/or 'Page X of Y'.
    """
    if not row:
        return False
    avg_y = sum(w['y_min'] for w in row) / len(row)
    row_text = " ".join([w['text'] for w in row]).lower()
    is_bottom = avg_y > 0.85  # Bottom 15% of page
    has_footer_pattern = 'powered by' in row_text or ('page' in row_text and 'of' in row_text)
    return is_bottom and has_footer_pattern


def group_rows_by_y_region(rows, images_info, y_threshold=0.08):
    """Group rows and images that share similar Y positions for side-by-side layout.

    Returns list of groups, where each group is:
    {
        'y_start': float,
        'y_end': float,
        'rows': list of rows in this Y region,
        'images': list of images in this Y region
    }
    """
    if not rows and not images_info:
        return []

    # Collect all content with Y positions
    content_items = []

    for row in rows:
        avg_y = sum(w['y_min'] for w in row) / len(row)
        min_x = min(w['x_min'] for w in row)
        max_x = max(w['x_max'] for w in row)
        content_items.append({
            'type': 'row',
            'y': avg_y,
            'x_min': min_x,
            'x_max': max_x,
            'data': row
        })

    for img in images_info:
        # Normalize image bbox to 0-1 range (assuming page dimensions)
        y_norm = img.get('y_norm', 0.1)  # Default to top area
        x_min = img.get('x_min_norm', 0)
        x_max = img.get('x_max_norm', 0.3)
        content_items.append({
            'type': 'image',
            'y': y_norm,
            'x_min': x_min,
            'x_max': x_max,
            'data': img
        })

    if not content_items:
        return []

    # Sort by Y position
    content_items.sort(key=lambda x: x['y'])

    # Group items with similar Y positions
    groups = []
    current_group = {
        'y_start': content_items[0]['y'],
        'y_end': content_items[0]['y'],
        'items': [content_items[0]]
    }

    for item in content_items[1:]:
        if abs(item['y'] - current_group['y_start']) < y_threshold:
            current_group['items'].append(item)
            current_group['y_end'] = max(current_group['y_end'], item['y'])
        else:
            groups.append(current_group)
            current_group = {
                'y_start': item['y'],
                'y_end': item['y'],
                'items': [item]
            }

    groups.append(current_group)

    return groups


def add_side_by_side_content(word_doc, group, page_images_data):
    """Add content that should be side-by-side using an invisible layout table.

    Args:
        word_doc: python-docx Document
        group: dict with 'items' list containing rows and/or images
        page_images_data: dict mapping image info to actual image bytes
    """
    items = group['items']

    # Sort items by X position (left to right)
    items.sort(key=lambda x: x['x_min'])

    # If only one item, just add it normally
    if len(items) == 1:
        item = items[0]
        if item['type'] == 'image':
            img_info = item['data']
            if 'bytes' in img_info:
                try:
                    img_stream = BytesIO(img_info['bytes'])
                    word_doc.add_picture(img_stream, width=Inches(1.5))
                except Exception as e:
                    logger.warning(f"Failed to add image: {e}")
        else:  # row
            row = item['data']
            row.sort(key=lambda w: w['x_min'])
            line_text = " ".join([w['text'] for w in row])
            if line_text.strip():
                para = word_doc.add_paragraph(line_text)
                para.paragraph_format.space_after = Pt(6)
        return

    # Multiple items - create layout table
    num_cols = len(items)
    table = word_doc.add_table(rows=1, cols=num_cols)

    # Make table invisible (no borders)
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            # Remove borders
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    # Add content to each cell
    for col_idx, item in enumerate(items):
        cell = table.rows[0].cells[col_idx]

        if item['type'] == 'image':
            img_info = item['data']
            if 'bytes' in img_info:
                try:
                    para = cell.paragraphs[0]
                    run = para.add_run()
                    img_stream = BytesIO(img_info['bytes'])
                    run.add_picture(img_stream, width=Inches(1.2))
                except Exception as e:
                    logger.warning(f"Failed to add image to cell: {e}")
        else:  # row
            row = item['data']
            row.sort(key=lambda w: w['x_min'])
            line_text = " ".join([w['text'] for w in row])
            if line_text.strip():
                cell.paragraphs[0].text = line_text

    # Add spacing after table
    word_doc.add_paragraph()


def convert_pdf_to_word_doctr(input_path: Path, output_path: Path) -> None:
    """Convert scanned PDF to Word using doctr OCR with layout preservation.

    Features:
    - Table detection based on column alignment
    - Side-by-side layout for header (logo + text)
    - Footer content placed in Word footer section
    """
    logger.info(f"Starting doctr conversion for: {input_path}")

    # Get doctr predictor
    predictor = get_doctr_predictor()

    # Open PDF and create Word document
    pdf_doc = fitz.open(str(input_path))
    word_doc = Document()

    # Set page size and margins
    for section in word_doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    # Track footer rows for the document (from first page)
    footer_rows_first_page = []

    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        logger.info(f"Processing page {page_num + 1}")

        # Get page dimensions for normalization
        page_rect = page.rect
        page_width = page_rect.width
        page_height = page_rect.height

        if page_num > 0:
            word_doc.add_page_break()

        # Convert page to high-res image for OCR
        scale = 2.0
        mat = fitz.Matrix(scale, scale)
        pix = page.get_pixmap(matrix=mat)

        # Save to temp file for doctr
        temp_img_path = OUTPUT_DIR / f"temp_page_{page_num}.png"
        pix.save(str(temp_img_path))

        # Run doctr OCR
        doc_file = DocumentFile.from_images(str(temp_img_path))
        result = predictor(doc_file)

        # Collect all words with their positions (normalized 0-1)
        all_words = []
        if result.pages:
            page_result = result.pages[0]
            for block in page_result.blocks:
                for line in block.lines:
                    for word in line.words:
                        if word.value.strip():
                            if hasattr(word, 'geometry') and word.geometry:
                                x_min, y_min = word.geometry[0]
                                x_max, y_max = word.geometry[1]
                                all_words.append({
                                    'text': word.value,
                                    'x_min': x_min,
                                    'y_min': y_min,
                                    'x_max': x_max,
                                    'y_max': y_max,
                                    'center_x': (x_min + x_max) / 2,
                                    'center_y': (y_min + y_max) / 2
                                })

        # Extract images with normalized positions
        page_images = extract_images_from_page(page)
        images_with_pos = []
        for img_info in page_images:
            x0, y0, x1, y1 = img_info['bbox']
            images_with_pos.append({
                'bytes': img_info['bytes'],
                'bbox': img_info['bbox'],
                'y_norm': y0 / page_height,
                'x_min_norm': x0 / page_width,
                'x_max_norm': x1 / page_width,
                'width': x1 - x0,
                'height': y1 - y0
            })

        # Group words into rows
        rows = detect_table_rows_from_words(all_words)
        logger.info(f"Found {len(rows)} text rows on page {page_num + 1}")

        # Separate footer content from regular content
        regular_rows = []
        footer_rows = []
        for row in rows:
            if is_footer_content(row):
                footer_rows.append(row)
                logger.info(f"Detected footer row: {' '.join([w['text'] for w in row])}")
            else:
                regular_rows.append(row)

        # Save footer rows (only from first page for processing)
        if footer_rows and page_num == 0:
            footer_rows_first_page = footer_rows.copy()

        # Add header images first (top 25% of page)
        header_images = [img for img in images_with_pos if img['y_norm'] < 0.25]
        for img_info in header_images:
            try:
                img_stream = BytesIO(img_info['bytes'])
                img_width_inches = min(img_info['width'] / 72, 2.0)
                word_doc.add_picture(img_stream, width=Inches(img_width_inches))
            except Exception as e:
                logger.warning(f"Failed to add header image: {e}")

        # Remove header images from further processing
        images_with_pos = [img for img in images_with_pos if img not in header_images]

        # Process remaining content with table detection
        table_start = None
        table_rows_collected = []
        non_table_rows = []

        for i, row in enumerate(regular_rows):
            row.sort(key=lambda w: w['x_min'])

            # Count significant gaps in this row
            gaps = []
            for j in range(len(row) - 1):
                gap = row[j + 1]['x_min'] - row[j]['x_max']
                if gap > 0.015:
                    gaps.append(gap)

            has_columns = len(gaps) >= 3  # At least 4 columns

            if has_columns:
                # Flush pending non-table rows
                if non_table_rows:
                    for ntr in non_table_rows:
                        ntr.sort(key=lambda w: w['x_min'])
                        line_text = " ".join([w['text'] for w in ntr])
                        if line_text.strip():
                            para = word_doc.add_paragraph(line_text)
                            para.paragraph_format.space_after = Pt(6)
                    non_table_rows = []

                if table_start is None:
                    table_start = i
                table_rows_collected.append(row)
            else:
                if table_rows_collected and len(table_rows_collected) >= 2:
                    _add_table_to_doc(word_doc, table_rows_collected)
                    table_rows_collected = []
                    table_start = None
                elif table_rows_collected:
                    non_table_rows.extend(table_rows_collected)
                    table_rows_collected = []
                    table_start = None
                non_table_rows.append(row)

        # Handle remaining table rows
        if table_rows_collected and len(table_rows_collected) >= 2:
            _add_table_to_doc(word_doc, table_rows_collected)
        elif table_rows_collected:
            non_table_rows.extend(table_rows_collected)

        # Add remaining non-table rows
        for row in non_table_rows:
            row.sort(key=lambda w: w['x_min'])
            line_text = " ".join([w['text'] for w in row])
            if line_text.strip():
                para = word_doc.add_paragraph(line_text)
                para.paragraph_format.space_after = Pt(6)

        # Clean up temp file
        try:
            temp_img_path.unlink()
        except Exception:
            pass

    # Add footer to document with proper left/right alignment
    if footer_rows_first_page:
        section = word_doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        # Separate footer words into left (x < 0.5) and right (x >= 0.5) parts
        all_footer_words = [w for row in footer_rows_first_page for w in row]
        left_words = sorted([w for w in all_footer_words if w['center_x'] < 0.5], key=lambda w: w['x_min'])
        right_words = sorted([w for w in all_footer_words if w['center_x'] >= 0.5], key=lambda w: w['x_min'])

        left_text = " ".join([w['text'] for w in left_words])  # "Powered by: CurerTech"
        right_text = " ".join([w['text'] for w in right_words])  # "Page 1 of 7"

        logger.info(f"Footer left: '{left_text}', right: '{right_text}'")

        # Create footer with tab stops for left and right alignment
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.clear()

        # Add left text
        run_left = footer_para.add_run(left_text)

        # Add tab to push to right
        footer_para.add_run("\t\t\t\t\t\t\t\t\t\t")

        # Add "Page " text
        footer_para.add_run("Page ")

        # Add PAGE field for current page number
        run_page = footer_para.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run_page._r.append(fld_char_begin)

        instr_text = OxmlElement('w:instrText')
        instr_text.text = "PAGE"
        run_page._r.append(instr_text)

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run_page._r.append(fld_char_end)

        # Add " of " text
        footer_para.add_run(" of ")

        # Add NUMPAGES field for total page count
        run_total = footer_para.add_run()
        fld_char_begin2 = OxmlElement('w:fldChar')
        fld_char_begin2.set(qn('w:fldCharType'), 'begin')
        run_total._r.append(fld_char_begin2)

        instr_text2 = OxmlElement('w:instrText')
        instr_text2.text = "NUMPAGES"
        run_total._r.append(instr_text2)

        fld_char_end2 = OxmlElement('w:fldChar')
        fld_char_end2.set(qn('w:fldCharType'), 'end')
        run_total._r.append(fld_char_end2)

        logger.info(f"Added footer with dynamic page numbers")

    pdf_doc.close()
    word_doc.save(str(output_path))
    logger.info(f"doctr conversion complete: {output_path}")


def _add_table_to_doc(word_doc, table_rows):
    """Add detected table rows to Word document.

    Args:
        word_doc: python-docx Document object
        table_rows: List of rows, each row is a list of word dicts
    """
    if not table_rows:
        return

    # Detect column boundaries from all words in table
    all_table_words = [w for row in table_rows for w in row]
    col_boundaries = detect_columns_from_words(all_table_words)

    if not col_boundaries or len(col_boundaries) < 3:
        # Fallback: estimate columns from first row
        first_row = table_rows[0]
        first_row.sort(key=lambda w: w['x_min'])

        # Find gaps in first row
        col_boundaries = [0.0]
        for j in range(len(first_row) - 1):
            gap_center = (first_row[j]['x_max'] + first_row[j + 1]['x_min']) / 2
            col_boundaries.append(gap_center)
        col_boundaries.append(1.0)

    num_cols = len(col_boundaries) - 1
    num_rows = len(table_rows)

    logger.info(f"Creating table with {num_rows} rows x {num_cols} cols")

    # Create Word table
    word_table = word_doc.add_table(rows=num_rows, cols=num_cols)
    word_table.style = 'Table Grid'
    word_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Fill table cells
    for row_idx, row_words in enumerate(table_rows):
        for word in row_words:
            # Find which column this word belongs to
            col_idx = -1
            for c in range(len(col_boundaries) - 1):
                if col_boundaries[c] <= word['center_x'] < col_boundaries[c + 1]:
                    col_idx = c
                    break

            # Handle edge case where word is at the right edge
            if col_idx == -1 and word['center_x'] >= col_boundaries[-2]:
                col_idx = num_cols - 1

            if 0 <= col_idx < num_cols:
                cell = word_table.rows[row_idx].cells[col_idx]
                current_text = cell.text.strip()
                if current_text:
                    cell.text = current_text + " " + word['text']
                else:
                    cell.text = word['text']

    word_doc.add_paragraph()  # Space after table

# Create uploads and outputs directories
UPLOAD_DIR = Path(tempfile.gettempdir()) / "converter_uploads"
OUTPUT_DIR = Path(tempfile.gettempdir()) / "converter_outputs"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


class ConversionError(Exception):
    """Custom exception for conversion errors."""

    pass


def get_output_extension(conversion_type: ConversionType) -> str:
    """Get the output file extension for a conversion type."""
    extensions = {
        ConversionType.PDF_TO_WORD: ".docx",
        ConversionType.WORD_TO_PDF: ".pdf",
        ConversionType.WORD_TO_EXCEL: ".xlsx",
        ConversionType.EXCEL_TO_WORD: ".docx",
    }
    return extensions[conversion_type]


def validate_input_file(filepath: Path, conversion_type: ConversionType) -> None:
    """Validate that the input file matches the expected format."""
    valid_extensions = {
        ConversionType.PDF_TO_WORD: [".pdf"],
        ConversionType.WORD_TO_PDF: [".docx", ".doc"],
        ConversionType.WORD_TO_EXCEL: [".docx", ".doc"],
        ConversionType.EXCEL_TO_WORD: [".xlsx", ".xls"],
    }

    ext = filepath.suffix.lower()
    if ext not in valid_extensions[conversion_type]:
        expected = ", ".join(valid_extensions[conversion_type])
        raise ConversionError(
            f"Invalid file type '{ext}' for {conversion_type.value}. Expected: {expected}"
        )


def is_image_based_pdf(pdf_doc) -> bool:
    """Check if PDF is image-based (scanned) with no extractable text."""
    total_text = ""
    for page_num in range(min(3, len(pdf_doc))):  # Check first 3 pages
        page = pdf_doc[page_num]
        total_text += page.get_text()
    return len(total_text.strip()) < 50  # Less than 50 chars = image-based


def detect_tables_opencv(img_array: np.ndarray, scale_factor: float = 1.0):
    """Detect tables in an image using OpenCV line detection.

    Returns list of table regions: [(x, y, w, h, rows, cols), ...]
    Each table includes detected row and column positions.
    """
    # Convert to grayscale
    if len(img_array.shape) == 3:
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
    else:
        gray = img_array

    # Threshold to get binary image
    _, binary = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)

    # Detect horizontal lines
    horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
    horizontal_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)

    # Detect vertical lines
    vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
    vertical_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, vertical_kernel, iterations=2)

    # Combine horizontal and vertical lines
    table_mask = cv2.add(horizontal_lines, vertical_lines)

    # Find contours (potential table regions)
    contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    tables = []
    min_table_area = 5000 * (scale_factor ** 2)  # Minimum area to be considered a table

    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        area = w * h

        if area > min_table_area and w > 100 * scale_factor and h > 50 * scale_factor:
            # Extract table region for row/column detection
            table_region = table_mask[y:y+h, x:x+w]

            # Find horizontal line positions (rows)
            h_proj = np.sum(table_region, axis=1)
            row_positions = np.where(h_proj > w * 0.3)[0]
            rows = _group_positions(row_positions, threshold=10)

            # Find vertical line positions (columns)
            v_proj = np.sum(table_region, axis=0)
            col_positions = np.where(v_proj > h * 0.3)[0]
            cols = _group_positions(col_positions, threshold=10)

            if len(rows) >= 2 and len(cols) >= 2:
                tables.append({
                    'bbox': (x, y, w, h),
                    'rows': [y + r for r in rows],
                    'cols': [x + c for c in cols]
                })

    return tables


def _group_positions(positions, threshold=10):
    """Group nearby positions into single line positions."""
    if len(positions) == 0:
        return []

    groups = []
    current_group = [positions[0]]

    for pos in positions[1:]:
        if pos - current_group[-1] <= threshold:
            current_group.append(pos)
        else:
            groups.append(int(np.mean(current_group)))
            current_group = [pos]

    if current_group:
        groups.append(int(np.mean(current_group)))

    return groups


def extract_images_from_page(page, min_size=50):
    """Extract images from a PDF page.

    Returns list of (image_bytes, bbox) tuples.
    """
    images = []
    image_list = page.get_images()

    for img_info in image_list:
        xref = img_info[0]
        try:
            base_image = page.parent.extract_image(xref)
            if base_image:
                img_bytes = base_image["image"]
                # Get image position on page
                for img_rect in page.get_image_rects(xref):
                    x0, y0, x1, y1 = img_rect
                    w, h = x1 - x0, y1 - y0
                    if w >= min_size and h >= min_size:
                        images.append({
                            'bytes': img_bytes,
                            'bbox': (x0, y0, x1, y1),
                            'ext': base_image.get("ext", "png")
                        })
        except Exception as e:
            logger.warning(f"Failed to extract image: {e}")
            continue

    return images


def convert_pdf_to_word_ocr(input_path: Path, output_path: Path) -> None:
    """Convert image-based PDF to Word using OCR with layout preservation.

    This enhanced version:
    - Extracts and embeds images/logos
    - Detects tables and recreates them in Word
    - OCRs text with layout-aware positioning
    """
    pdf_doc = fitz.open(str(input_path))
    word_doc = Document()
    reader = get_ocr_reader()

    # Set page dimensions (A4)
    section = word_doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        logger.info(f"Processing page {page_num + 1}")

        if page_num > 0:
            word_doc.add_page_break()

        # Get page dimensions
        page_rect = page.rect
        page_width = page_rect.width
        page_height = page_rect.height

        # Convert page to image for OCR (2x scale for better quality)
        scale_factor = 2.0
        mat = fitz.Matrix(scale_factor, scale_factor)
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        img = Image.open(BytesIO(img_data))
        img_array = np.array(img)

        # 1. Extract embedded images/logos from PDF
        page_images = extract_images_from_page(page)
        logger.info(f"Found {len(page_images)} embedded images")

        # 2. Detect tables using OpenCV
        tables = detect_tables_opencv(img_array, scale_factor)
        logger.info(f"Detected {len(tables)} tables")

        # Create table bounding boxes (scaled back to page coords)
        table_regions = []
        for table in tables:
            x, y, w, h = table['bbox']
            table_regions.append((
                x / scale_factor,
                y / scale_factor,
                (x + w) / scale_factor,
                (y + h) / scale_factor
            ))

        # 3. Run OCR on entire page
        ocr_results = reader.readtext(img_array)

        # Separate OCR results into table cells and free text
        table_ocr = {i: [] for i in range(len(tables))}
        free_text_ocr = []

        for (bbox, text, conf) in ocr_results:
            # Get center point of text (scaled to page coords)
            center_x = (bbox[0][0] + bbox[2][0]) / 2 / scale_factor
            center_y = (bbox[0][1] + bbox[2][1]) / 2 / scale_factor

            in_table = False
            for i, (tx1, ty1, tx2, ty2) in enumerate(table_regions):
                if tx1 <= center_x <= tx2 and ty1 <= center_y <= ty2:
                    table_ocr[i].append((bbox, text, conf))
                    in_table = True
                    break

            if not in_table:
                free_text_ocr.append((bbox, text, conf, center_y))

        # 4. Add embedded images to document (at top)
        for img_info in page_images:
            try:
                img_stream = BytesIO(img_info['bytes'])
                # Calculate width in inches (max 6 inches wide)
                x0, y0, x1, y1 = img_info['bbox']
                img_width = min((x1 - x0) / 72, 6.0)  # Convert points to inches
                word_doc.add_picture(img_stream, width=Inches(img_width))
            except Exception as e:
                logger.warning(f"Failed to add image: {e}")

        # 5. Add detected tables
        for i, table in enumerate(tables):
            rows = table['rows']
            cols = table['cols']

            if len(rows) < 2 or len(cols) < 2:
                continue

            num_rows = len(rows) - 1
            num_cols = len(cols) - 1

            logger.info(f"Creating table with {num_rows} rows x {num_cols} cols")

            word_table = word_doc.add_table(rows=num_rows, cols=num_cols)
            word_table.style = 'Table Grid'
            word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Fill table cells with OCR text
            for (bbox, text, conf) in table_ocr[i]:
                # Find which cell this text belongs to
                center_x = (bbox[0][0] + bbox[2][0]) / 2 / scale_factor
                center_y = (bbox[0][1] + bbox[2][1]) / 2 / scale_factor

                # Find row and column
                row_idx = -1
                col_idx = -1

                for r in range(len(rows) - 1):
                    if rows[r] / scale_factor <= center_y <= rows[r + 1] / scale_factor:
                        row_idx = r
                        break

                for c in range(len(cols) - 1):
                    if cols[c] / scale_factor <= center_x <= cols[c + 1] / scale_factor:
                        col_idx = c
                        break

                if 0 <= row_idx < num_rows and 0 <= col_idx < num_cols:
                    cell = word_table.rows[row_idx].cells[col_idx]
                    if cell.text:
                        cell.text += " " + text
                    else:
                        cell.text = text

            word_doc.add_paragraph()  # Space after table

        # 6. Add free text (outside tables) sorted by Y position
        free_text_ocr.sort(key=lambda x: x[3])  # Sort by Y position

        # Group into lines
        current_y = -100
        current_line = []
        lines = []
        line_threshold = 20 / scale_factor

        for (bbox, text, conf, y_pos) in free_text_ocr:
            if abs(y_pos - current_y) > line_threshold:
                if current_line:
                    lines.append(current_line)
                current_line = [(bbox, text)]
                current_y = y_pos
            else:
                current_line.append((bbox, text))

        if current_line:
            lines.append(current_line)

        # Output lines
        for line_items in lines:
            # Sort by X position within line
            line_items.sort(key=lambda x: x[0][0][0])
            line_text = " ".join([item[1] for item in line_items])

            if line_text.strip():
                para = word_doc.add_paragraph()
                run = para.add_run(line_text)
                run.font.size = Pt(11)

    pdf_doc.close()
    word_doc.save(str(output_path))
    logger.info(f"OCR conversion complete: {output_path}")


def convert_pdf_to_word_text(input_path: Path, output_path: Path) -> None:
    """Convert text-based PDF to Word using PyMuPDF extraction."""
    pdf_doc = fitz.open(str(input_path))
    word_doc = Document()

    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]

        if page_num > 0:
            word_doc.add_page_break()

        # Extract tables first
        table_finder = page.find_tables()
        tables = table_finder.tables if table_finder else []
        table_rects = [table.bbox for table in tables]

        # Extract and add tables
        for table in tables:
            data = table.extract()
            if data and len(data) > 0 and len(data[0]) > 0:
                rows = len(data)
                cols = len(data[0])
                word_table = word_doc.add_table(rows=rows, cols=cols)
                word_table.style = 'Table Grid'

                for i, row_data in enumerate(data):
                    for j, cell_text in enumerate(row_data):
                        if cell_text:
                            word_table.rows[i].cells[j].text = str(cell_text)

                word_doc.add_paragraph()

        # Extract text blocks
        blocks = page.get_text("dict")["blocks"]

        for block in blocks:
            if block["type"] == 0:
                block_rect = fitz.Rect(block["bbox"])

                in_table = any(
                    block_rect.intersects(fitz.Rect(tr)) for tr in table_rects
                )
                if in_table:
                    continue

                for line in block.get("lines", []):
                    line_text = ""
                    max_size = 11
                    is_bold = False

                    for span in line.get("spans", []):
                        line_text += span.get("text", "")
                        max_size = max(max_size, span.get("size", 11))
                        if "bold" in span.get("font", "").lower():
                            is_bold = True

                    line_text = line_text.strip()
                    if line_text:
                        para = word_doc.add_paragraph()
                        run = para.add_run(line_text)

                        if max_size > 16:
                            run.bold = True
                            run.font.size = Pt(14)
                        elif max_size > 13:
                            run.bold = True
                            run.font.size = Pt(12)
                        else:
                            run.font.size = Pt(11)
                            if is_bold:
                                run.bold = True

    pdf_doc.close()
    word_doc.save(str(output_path))


def convert_pdf_to_word_pdf2docx(input_path: Path, output_path: Path) -> None:
    """Convert PDF to Word using pdf2docx for high-fidelity layout preservation.

    This preserves tables, images, text formatting, and layout positioning.
    """
    logger.info("Using pdf2docx for high-fidelity conversion")
    cv = Pdf2DocxConverter(str(input_path))
    cv.convert(str(output_path))
    cv.close()


def convert_pdf_to_word(input_path: Path, output_path: Path) -> None:
    """Convert PDF to Word document.

    Uses pdf2docx for high-fidelity conversion that preserves:
    - Tables
    - Images and logos
    - Text formatting
    - Layout positioning

    Falls back to OCR for image-based (scanned) PDFs.

    Args:
        input_path: Path to the input PDF file.
        output_path: Path where the output DOCX will be saved.

    Raises:
        ConversionError: If conversion fails.
    """
    try:
        logger.info(f"convert_pdf_to_word called with input: {input_path}")
        pdf_doc = fitz.open(str(input_path))

        is_image_based = is_image_based_pdf(pdf_doc)
        logger.info(f"is_image_based_pdf result: {is_image_based}")
        pdf_doc.close()

        if is_image_based:
            # For scanned/image-based PDFs, use doctr for OCR
            logger.info("Using doctr conversion for image-based PDF")
            convert_pdf_to_word_doctr(input_path, output_path)
        else:
            # For text-based PDFs, use pdf2docx which preserves original formatting better
            logger.info("Using pdf2docx for text-based PDF")
            convert_pdf_to_word_pdf2docx(input_path, output_path)

        logger.info(f"Conversion complete, output: {output_path}")

    except Exception as e:
        logger.error(f"PDF to Word conversion failed: {e}")
        raise ConversionError(f"PDF to Word conversion failed: {e}") from e


def convert_word_to_pdf(input_path: Path, output_path: Path) -> None:
    """Convert Word document to PDF.

    Note: This requires LibreOffice for high-quality conversion.
    Falls back to a basic placeholder if LibreOffice is not available.

    Args:
        input_path: Path to the input DOCX file.
        output_path: Path where the output PDF will be saved.

    Raises:
        ConversionError: If conversion fails.
    """
    try:
        # Try using LibreOffice if available
        import subprocess

        # Check if LibreOffice is installed
        libreoffice_paths = [
            "libreoffice",
            "soffice",
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
            "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
        ]

        libreoffice_cmd = None
        for path in libreoffice_paths:
            try:
                result = subprocess.run(
                    [path, "--version"],
                    capture_output=True,
                    timeout=5,
                )
                if result.returncode == 0:
                    libreoffice_cmd = path
                    break
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue

        if libreoffice_cmd:
            # Use LibreOffice for conversion
            subprocess.run(
                [
                    libreoffice_cmd,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_path.parent),
                    str(input_path),
                ],
                check=True,
                capture_output=True,
                timeout=120,
            )
            # LibreOffice creates file with same name but .pdf extension
            temp_output = output_path.parent / (input_path.stem + ".pdf")
            if temp_output != output_path and temp_output.exists():
                temp_output.rename(output_path)
        else:
            # Fallback: Create a simple PDF with python-docx content
            raise ConversionError(
                "LibreOffice not found. Please install LibreOffice for Word to PDF conversion. "
                "Download from: https://www.libreoffice.org/download/download/"
            )
    except subprocess.CalledProcessError as e:
        raise ConversionError(f"Word to PDF conversion failed: {e.stderr.decode()}") from e
    except Exception as e:
        if isinstance(e, ConversionError):
            raise
        raise ConversionError(f"Word to PDF conversion failed: {e}") from e


def convert_word_to_excel(input_path: Path, output_path: Path) -> None:
    """Convert Word document tables to Excel spreadsheet.

    Extracts all tables from the Word document and saves them to Excel.

    Args:
        input_path: Path to the input DOCX file.
        output_path: Path where the output XLSX will be saved.

    Raises:
        ConversionError: If conversion fails.
    """
    try:
        doc = Document(str(input_path))
        wb = Workbook()

        # Remove default sheet if we have tables
        if doc.tables:
            wb.remove(wb.active)

        if not doc.tables:
            # If no tables, extract paragraphs as single column
            ws = wb.active
            ws.title = "Content"
            for i, para in enumerate(doc.paragraphs, 1):
                if para.text.strip():
                    ws.cell(row=i, column=1, value=para.text)
        else:
            # Extract each table to a separate sheet
            for idx, table in enumerate(doc.tables, 1):
                ws = wb.create_sheet(title=f"Table_{idx}")
                for i, row in enumerate(table.rows, 1):
                    for j, cell in enumerate(row.cells, 1):
                        ws.cell(row=i, column=j, value=cell.text)

        wb.save(str(output_path))
    except Exception as e:
        raise ConversionError(f"Word to Excel conversion failed: {e}") from e


def convert_excel_to_word(input_path: Path, output_path: Path) -> None:
    """Convert Excel spreadsheet to Word document with tables.

    Args:
        input_path: Path to the input XLSX file.
        output_path: Path where the output DOCX will be saved.

    Raises:
        ConversionError: If conversion fails.
    """
    try:
        wb = load_workbook(str(input_path))
        doc = Document()

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # Add sheet name as heading
            doc.add_heading(sheet_name, level=1)

            # Get dimensions
            max_row = ws.max_row
            max_col = ws.max_column

            if max_row == 0 or max_col == 0:
                doc.add_paragraph("(Empty sheet)")
                continue

            # Create table in Word
            table = doc.add_table(rows=max_row, cols=max_col)
            table.style = "Table Grid"

            # Fill table with data
            for i, row in enumerate(ws.iter_rows(min_row=1, max_row=max_row)):
                for j, cell in enumerate(row):
                    value = cell.value if cell.value is not None else ""
                    table.rows[i].cells[j].text = str(value)

            doc.add_paragraph()  # Add spacing between tables

        doc.save(str(output_path))
    except Exception as e:
        raise ConversionError(f"Excel to Word conversion failed: {e}") from e


def convert_file(
    input_path: Path,
    conversion_type: ConversionType,
    output_filename: str | None = None,
) -> Path:
    """Convert a file based on the specified conversion type.

    Args:
        input_path: Path to the input file.
        conversion_type: Type of conversion to perform.
        output_filename: Optional custom output filename.

    Returns:
        Path to the converted file.

    Raises:
        ConversionError: If conversion fails or file type is invalid.
    """
    # Validate input file
    validate_input_file(input_path, conversion_type)

    # Generate output path
    if output_filename:
        output_path = OUTPUT_DIR / output_filename
    else:
        output_ext = get_output_extension(conversion_type)
        output_path = OUTPUT_DIR / (input_path.stem + "_converted" + output_ext)

    # Perform conversion
    converters = {
        ConversionType.PDF_TO_WORD: convert_pdf_to_word,
        ConversionType.WORD_TO_PDF: convert_word_to_pdf,
        ConversionType.WORD_TO_EXCEL: convert_word_to_excel,
        ConversionType.EXCEL_TO_WORD: convert_excel_to_word,
    }

    converter_func = converters[conversion_type]
    converter_func(input_path, output_path)

    if not output_path.exists():
        raise ConversionError("Conversion completed but output file not found")

    return output_path


def cleanup_file(filepath: Path) -> None:
    """Remove a temporary file."""
    try:
        if filepath.exists():
            os.remove(filepath)
    except OSError:
        pass  # Ignore cleanup errors
